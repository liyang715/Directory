import os
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from docx import Document

def read_docx_tables(file_path):
    doc = Document(file_path)
    tables = [t for t in doc.tables]
    return tables

def extract_names_from_tables_v2(tables, keyword):
    names = []
    for table in tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            if keyword in row_text:
                name_index = row_text.index(keyword) + 2  # Adjust the index to the second cell after "姓 名"
                if name_index < len(row_text):
                    name = row_text[name_index]
                    if name:  # Ensure the name is not empty or just whitespaces
                        names.append(name)
    return names

def extract_odd_pages(names):
    return names[::2]  # Select every other name starting from the first one

def create_index(names):
    return [(i*2+1, name) for i, name in enumerate(names)]

def write_index_to_docx_v2(index, file_path):
    doc = Document()
    for i, (page, name) in enumerate(index, start=1):
        entry = f"{i}.{name}" + "."*(30-len(name)) + str(page)  # Add dots between the name and the page number
        doc.add_paragraph(entry)
    doc.save(file_path)

# Create a file dialog to select the input docx file
Tk().withdraw()  # Close the blank tkinter window
input_file_path = askopenfilename(filetypes=[('Word Document', '*.docx')])  # Show the file dialog

# Read the tables from the docx file
tables = read_docx_tables(input_file_path)

# Extract the names from the tables
keyword = "姓  名"
names = extract_names_from_tables_v2(tables, keyword)

# Select the names from odd pages
names_odd_pages = extract_odd_pages(names)

# Create the index
index = create_index(names_odd_pages)

# Set the output file path to be in the same directory as the input file, with the name "目录.docx"
output_file_path = os.path.join(os.path.dirname(input_file_path), "干审表目录.docx")

# Write the index to the output docx file
write_index_to_docx_v2(index, output_file_path)

